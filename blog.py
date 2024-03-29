# Import flash from Flask
from flask import Flask, render_template, request, redirect, url_for, flash
from flask_sqlalchemy import SQLAlchemy
from flask import send_from_directory
from flask import send_file
from datetime import datetime
from openpyxl import Workbook 
from flask_login import current_user
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user
from docx import Document
from PIL import Image
from sqlalchemy import or_
from docx.shared import Inches
import io
import os

app = Flask(__name__, static_url_path='/static')

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///blogs.db'  # Change 'your_database.db' to your desired database filename
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)
login_manager = LoginManager(app)
# Define database models
# User model
class User(db.Model, UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)

    # Implement the required methods for Flask-Login
    def is_authenticated(self):
        return True

    def is_active(self):
        return True

    def is_anonymous(self):
        return False

    def get_id(self):
        return str(self.id)

class Post(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    content = db.Column(db.Text, nullable=False)
    image = db.Column(db.LargeBinary, nullable=True)  # Ensure this line is present
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    author_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

    author = db.relationship('User', backref=db.backref('posts', lazy=True))



# Routes
@app.route('/')
def home():
    if not current_user.is_authenticated:
        return redirect(url_for('login'))
    posts = Post.query.all()
    return render_template('index.html', posts=posts)


@app.route('/post_image/<int:post_id>')
def post_image(post_id):
    post = Post.query.get_or_404(post_id)
    if post.image:
        # Open the image using Pillow to determine its format
        image = Image.open(io.BytesIO(post.image))
        # Convert the image to bytes
        img_byte_array = io.BytesIO()
        image.save(img_byte_array, format=image.format)
        img_byte_array.seek(0)
        # Return the image data with the appropriate MIME type
        return send_file(img_byte_array, mimetype=f'image/{image.format.lower()}')
    else:
        # If there's no image data, return a placeholder image or an error message
        return "No image available"
    

# Login view
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username, password=password).first()
        if user:
            login_user(user)
            return redirect(url_for('home'))
        else:
            return "Invalid username or password"
    return render_template('login.html')

@app.route('/search', methods=['GET'])
def search():
    query = request.args.get('query')
    if query:
        # Search for posts containing the query in title or content
        search_results = Post.query.filter(or_(Post.title.contains(query), Post.content.contains(query))).all()
    else:
        # If no query is provided, show all posts
        search_results = Post.query.all()
    return render_template('search_results.html', search_results=search_results, query=query)

# Protected route
@app.route('/protected')
@login_required
def protected():
    return "This page is protected. Only authenticated users can access it."

# User loader function
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User(username=username, password=password)
        db.session.add(user)
        db.session.commit()
        return redirect(url_for('login'))
    return render_template('signup.html')

from flask import redirect, url_for

@app.route('/add_post', methods=['GET', 'POST'])
def add_post():
    if request.method == 'POST':
        # Process form submission
        title = request.form['title']
        content = request.form['content']
        image = request.files['image'].read()  # Read image data from form 
        if current_user.is_authenticated:  # Check if user is logged in
            # Associate the post with the current user
            post = Post(title=title, content=content, image=image, author_id=current_user.id)
            db.session.add(post)
            db.session.commit()
            return redirect(url_for('home'))  # Redirect to home page after adding post
        else:
            return redirect(url_for('login'))  # Redirect to login page if user is not logged in
    else:
        # Render add post form
        if current_user.is_authenticated:  # Check if user is logged in
            return render_template('add_post.html')
        else:
            return redirect(url_for('login'))  # Redirect to login page if user is not logged in


# Modify your delete_post route
@app.route('/delete_post/<int:id>', methods=['POST'])
@login_required  # Ensure that only authenticated users can delete posts
def delete_post(id):
    post = Post.query.get_or_404(id)
    # Check if the current user is the author of the post
    if current_user.is_authenticated and current_user.id == post.author_id:
        db.session.delete(post)
        db.session.commit()
        flash('Post deleted successfully', 'success')
    else:
        flash('You are not authorized to delete this post', 'error')
    return redirect(url_for('home'))

@app.route('/delete_all_posts', methods=['POST'])
@login_required  # Ensure that only authenticated users can delete all posts
def delete_all_posts():
    # Check if the current user is authorized to delete all posts (optional)
    if current_user.is_authenticated:
        # Delete all posts from the database
        Post.query.delete()
        db.session.commit()
        flash('All posts deleted successfully', 'success')  # Display success message
    else:
        flash('You are not authorized to delete all posts', 'error')  # Display error message
    return redirect(url_for('home'))

@app.route('/delete_all_users', methods=['POST'])
@login_required  # Ensure that only authenticated users can delete all users
def delete_all_users():
    # Check if the current user is authorized to delete all users (optional)
    if current_user.is_authenticated:
        # Delete all users from the database
        User.query.delete()
        db.session.commit()
        flash('All users deleted successfully', 'success')  # Display success message
    else:
        flash('You are not authorized to delete all users', 'error')  # Display error message
    return redirect(url_for('home'))

# Flask route to handle updating a post
@app.route('/update_post/<int:post_id>', methods=['GET', 'POST'])
@login_required
def update_post(post_id):
    post = Post.query.get_or_404(post_id)
    if request.method == 'POST':
        # Update post title and content
        post.title = request.form['title']
        post.content = request.form['content']
        # Update image if provided
        if 'image' in request.files:
            image_file = request.files['image']
            if image_file:
                post.image = image_file.read()
        db.session.commit()
        return redirect(url_for('home'))
    return render_template('update_post.html', post=post)

def generate_word_document(post):
    # Create a new Word document
    document = Document()
    # Add the title to the document
    document.add_heading(post.title, level=1)
    # Add the image to the document, if available
    if post.image:
        image_stream = io.BytesIO(post.image)
        document.add_picture(image_stream, width=Inches(2))
    # Add the content to the document
    document.add_paragraph(post.content)
    # Add the author's name to the document
    document.add_paragraph(f"Author: {post.author.username}")
    return document


@app.route('/save_as_word/<int:id>')
def save_as_word(id):
    post = Post.query.get_or_404(id)
    # Generate the Word document
    document = generate_word_document(post)
    # Get the path to the user's Downloads directory
    downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads')
    filename = f"{post.title}.docx"
    file_path = os.path.join(downloads_path, filename)
    document.save(file_path)
    # Return the Word document as a file attachment
    return send_file(file_path, as_attachment=True)

@app.route('/post/<int:post_id>')
def view_post(post_id):
    post = Post.query.get_or_404(post_id)
    return render_template('view_post.html', post=post)


@app.route('/logout', methods=['POST'])
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

if __name__ == '__main__':
    # Create the database tables within the application context
    with app.app_context():
        db.create_all()
    app.run(debug=True, port=5000)
